from flask import Flask, request, send_file, render_template, jsonify
from flask import redirect, url_for
from docx import Document
from docx.shared import RGBColor, Pt
from reportlab.platypus import SimpleDocTemplate, Paragraph, Preformatted
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from html.parser import HTMLParser
import io
from citation import Citation
import json
import os

app = Flask(__name__)

DATA_FILE = 'data.json'

# Global storage
citations = []
content = ''
raw_tex = ''  # original .tex source preserved for round-trip export

# keep a running counter to assign unique citation IDs
next_cid = 1

# remember current sort settings (None means unsorted)
sort_by = None
sort_dir = 'asc'  # 'asc' or 'desc'

def load_data():
    global citations, content, next_cid, sort_by, sort_dir, raw_tex
    if os.path.exists(DATA_FILE):
        try:
            with open(DATA_FILE, 'r') as f:
                data = json.load(f)
                citations = [Citation(**c) for c in data.get('citations', [])]
                content = data.get('content', '')
                next_cid = data.get('next_cid', 1)
                sort_by = data.get('sort_by', None)
                sort_dir = data.get('sort_dir', 'asc')
                raw_tex  = data.get('raw_tex', '')
                print(f"[DEBUG] Loaded content length: {len(content)} chars")
                if len(content) > 0:
                    print(f"[DEBUG] Content preview: {content[:100]}...")
        except Exception as e:
            print(f"[DEBUG] Error loading data: {e}")
            pass  # if load fails, keep defaults

def save_data():
    data = {
        'citations': [{'cid': c.cid, 'author': c.author, 'year': c.year, 'source_type': c.source_type, 'link': c.link, 'pages': c.pages, 'volume': c.volume, 'issue': c.issue, 'publisher': c.publisher, 'doi': c.doi, 'title': c.title, 'style': c.style, 'inlineCites': c.inlineCites} for c in citations],
        'content': content,
        'raw_tex': raw_tex,
        'next_cid': next_cid,
        'sort_by': sort_by,
        'sort_dir': sort_dir
    }
    try:
        with open(DATA_FILE, 'w') as f:
            json.dump(data, f)
    except Exception as e:
        print(f"[DEBUG] Error saving data: {e}")


def docx_to_html(doc):
    """
    Convert a python-docx Document to HTML with formatting preserved.
    Handles bold, italic, underline, font size, font name, color, and paragraphs.
    """
    html_parts = []
    
    for para in doc.paragraphs:
        # Skip completely empty paragraphs but preserve them as empty <p>
        if not para.text.strip():
            html_parts.append('<p><br></p>')
            continue
        
        para_html = '<p'
        
        # Handle paragraph alignment
        alignment = para.alignment
        if alignment:
            align_map = {0: 'left', 1: 'center', 2: 'right', 3: 'justify'}
            align_value = align_map.get(alignment, 'left')
            para_html += f' style="text-align: {align_value};"'
        
        para_html += '>'
        
        for run in para.runs:
            text = run.text
            if not text:
                continue
            
            # Build inline styles
            styles = []
            
            # Font size
            if run.font.size:
                font_size_pt = run.font.size.pt
                styles.append(f'font-size: {int(font_size_pt)}pt')
            
            # Font name
            if run.font.name:
                # Escape font name for CSS
                font_name = run.font.name.replace('"', "'")
                styles.append(f"font-family: '{font_name}'")
            
            # Font color
            if run.font.color and run.font.color.rgb:
                color = str(run.font.color.rgb)
                if not color.startswith('#'):
                    color = f'#{color}'
                styles.append(f'color: {color}')
            
            # Build the span or formatting tags
            inner_html = text
            
            # Apply formatting tags
            if run.font.underline:
                inner_html = f'<u>{inner_html}</u>'
            if run.font.italic:
                inner_html = f'<i>{inner_html}</i>'
            if run.font.bold:
                inner_html = f'<b>{inner_html}</b>'
            
            # Wrap with span if we have styles
            if styles:
                style_string = ';'.join(styles)
                inner_html = f'<span style="{style_string}">{inner_html}</span>'
            
            para_html += inner_html
        
        para_html += '</p>'
        html_parts.append(para_html)
    
    return ''.join(html_parts)


def html_to_plain_text(html):
    """
    Convert HTML to plain text by stripping tags.
    """
    class HTMLStripper(HTMLParser):
        def __init__(self):
            super().__init__()
            self.reset()
            self.strict = False
            self.convert_charrefs = True
            self.text = []
        
        def handle_data(self, data):
            self.text.append(data)
        
        def get_data(self):
            return ''.join(self.text)
    
    stripper = HTMLStripper()
    stripper.feed(html)
    return stripper.get_data()


# ---------------- HOME ----------------
@app.route('/')
def index():
    load_data()
    # apply existing sort before rendering
    global citations
    if sort_by:
        citations.sort(
            key=lambda c: (getattr(c, sort_by) or "").lower(),
            reverse=(sort_dir == 'desc')
        )
    
    return render_template(
        "page.html",
        citations=citations,
        content=content,
        sort_by=sort_by,
        sort_dir=sort_dir
    )

# ---------------- LOAD FILE ----------------
@app.route('/load', methods=['POST'])
def load_file():
    global content
    global citations
    global sort_by
    global sort_dir
    global raw_tex

    file = request.files.get('file')
    print(f"[DEBUG] Request files: {request.files}")
    print(f"[DEBUG] File object: {file}")

    if not file or not file.filename:
        print("[DEBUG] No file provided in request")
        return redirect(url_for('index'))

    filename = file.filename.lower()
    print(f"[DEBUG] Loading file: {filename}")

    file_processed = False  # Track if file was successfully processed

    try:
        import html as _html

        if filename.endswith('.txt'):
            text_content = file.read().decode('utf-8')
            paragraphs = []
            for line in text_content.split('\n'):
                if line.strip():
                    paragraphs.append(f'<p>{_html.escape(line)}</p>')
                else:
                    paragraphs.append('<p><br></p>')
            content = ''.join(paragraphs)
            file_processed = True
            print(f"[DEBUG] TXT imported successfully, content length: {len(content)} chars")

        elif filename.endswith('.docx'):
            doc = Document(file)
            content = docx_to_html(doc)
            file_processed = True
            print(f"[DEBUG] DOCX imported successfully, content length: {len(content)} chars")

        elif filename.endswith('.pdf'):
            import pdfplumber
            import io as _io
            pdf_bytes = file.read()
            print(f"[DEBUG] PDF bytes read: {len(pdf_bytes)}")
            paragraphs = []
            with pdfplumber.open(_io.BytesIO(pdf_bytes)) as pdf:
                print(f"[DEBUG] PDF pages: {len(pdf.pages)}")
                for page in pdf.pages:
                    page_text = page.extract_text() or ''
                    print(f"[DEBUG] Page text length: {len(page_text)}")
                    for line in page_text.split('\n'):
                        if line.strip():
                            paragraphs.append(f'<p>{_html.escape(line)}</p>')
                        else:
                            paragraphs.append('<p><br></p>')
                    paragraphs.append('<p><br></p>')
            content = ''.join(paragraphs)
            file_processed = True
            print(f"[DEBUG] PDF imported successfully, content length: {len(content)} chars")

        elif filename.endswith('.tex'):
            import re as _re2
            raw_src = file.read().decode('utf-8', errors='replace')
            raw_tex = raw_src  # preserve for round-trip export

            # Convert LaTeX to rich HTML for the editor
            def tex_to_html(t):
                # Strip line comments (but keep the newline so paragraph breaks survive)
                t = _re2.sub(r'%[^\n]*', '', t)

                # Extract only the document body
                doc_match = _re2.search(r'\\begin\{document\}', t)
                if doc_match:
                    t = t[doc_match.end():]
                end_doc_match = _re2.search(r'\\end\{document\}', t)
                if end_doc_match:
                    t = t[:end_doc_match.start()]

                # Drop whole skipped environments (non-greedy, DOTALL)
                for env in ('figure', 'table', 'lstlisting', 'verbatim', 'tikzpicture'):
                    t = _re2.sub(
                        r'\\begin\{' + env + r'\}.*?\\end\{' + env + r'\}',
                        '', t, flags=_re2.DOTALL
                    )

                # Convert list environments to HTML (repeat for nesting)
                def convert_list_env(m):
                    env = m.group(1)
                    inner = m.group(2)
                    tag = 'ol' if env == 'enumerate' else 'ul'
                    items = _re2.split(r'\\item\b', inner)
                    li_parts = ''.join(
                        f'<li>{item.strip()}</li>'
                        for item in items if item.strip()
                    )
                    return f'\n<{tag}>{li_parts}</{tag}>\n'
                for _ in range(6):
                    new_t = _re2.sub(
                        r'\\begin\{(enumerate|itemize)\}(.*?)\\end\{\1\}',
                        convert_list_env, t, flags=_re2.DOTALL
                    )
                    if new_t == t:
                        break
                    t = new_t

                # quote/quotation/abstract -> blockquote
                t = _re2.sub(
                    r'\\begin\{(quote|quotation|abstract)\}(.*?)\\end\{\1\}',
                    lambda m: f'\n<blockquote>{m.group(2).strip()}</blockquote>\n',
                    t, flags=_re2.DOTALL
                )

                # Headings — add double newlines so they become their own chunks
                t = _re2.sub(r'\\chapter\*?\{([^}]*)\}',      r'\n\n<h1>\1</h1>\n\n', t)
                t = _re2.sub(r'\\section\*?\{([^}]*)\}',       r'\n\n<h2>\1</h2>\n\n', t)
                t = _re2.sub(r'\\subsection\*?\{([^}]*)\}',    r'\n\n<h3>\1</h3>\n\n', t)
                t = _re2.sub(r'\\subsubsection\*?\{([^}]*)\}', r'\n\n<h4>\1</h4>\n\n', t)
                t = _re2.sub(r'\\paragraph\*?\{([^}]*)\}',     r'\n\n<h5>\1</h5>\n\n', t)

                # Silently drop standalone control words (no content lost)
                for cmd in ('maketitle', 'tableofcontents', 'newpage', 'clearpage',
                            'cleardoublepage', 'noindent', 'medskip', 'bigskip',
                            'smallskip', 'hline', 'toprule', 'midrule', 'bottomrule'):
                    t = t.replace('\\' + cmd, '')

                # Inline formatting — only match single-line {arg} to avoid eating paragraphs
                t = _re2.sub(r'\\textbf\{([^}]*)\}',    r'<strong>\1</strong>', t)
                t = _re2.sub(r'\\textit\{([^}]*)\}',    r'<em>\1</em>', t)
                t = _re2.sub(r'\\emph\{([^}]*)\}',      r'<em>\1</em>', t)
                t = _re2.sub(r'\\underline\{([^}]*)\}', r'<u>\1</u>', t)
                t = _re2.sub(r'\\texttt\{([^}]*)\}',    r'<code>\1</code>', t)
                t = _re2.sub(r'\\footnote\{([^}]*)\}',  r' (\1)', t)
                t = _re2.sub(r'\\href\{[^}]*\}\{([^}]*)\}', r'\1', t)
                t = _re2.sub(r'\\url\{([^}]*)\}',       r'\1', t)

                # Drop citation/ref commands entirely
                t = _re2.sub(r'\\(cite|citep|citet|citealt|ref|eqref|label|bibitem)\{[^}]*\}', '', t)

                # LaTeX line-break -> actual newline
                t = t.replace('\\\\', '\n')

                # Drop all remaining \begin{...} and \end{...} tags (environment wrappers)
                t = _re2.sub(r'\\begin\{[^}]*\}', '', t)
                t = _re2.sub(r'\\end\{[^}]*\}', '', t)

                # Drop remaining commands that take a single {arg} — strip the command
                # but KEEP the argument text so prose isn't lost
                t = _re2.sub(r'\\[a-zA-Z]+\*?\{([^}]*)\}', r'\1', t)

                # Drop bare commands (no argument) — these are always markup, never prose
                t = _re2.sub(r'\\[a-zA-Z]+\*?', '', t)

                # Remove any leftover stray braces
                t = t.replace('{', '').replace('}', '')

                # Build HTML paragraphs.
                # Split on 2+ newlines for paragraph breaks.
                # Single newlines within a paragraph are joined as spaces.
                BLOCK_RE = _re2.compile(r'^\s*<(h[1-6]|ul|ol|blockquote)')
                out = []
                for chunk in _re2.split(r'\n{2,}', t):
                    lines = [l.strip() for l in chunk.split('\n') if l.strip()]
                    if not lines:
                        continue
                    text_buf = []
                    for line in lines:
                        if BLOCK_RE.match(line):
                            if text_buf:
                                out.append(f'<p>{" ".join(text_buf)}</p>')
                                text_buf = []
                            out.append(line)
                        else:
                            text_buf.append(line)
                    if text_buf:
                        out.append(f'<p>{" ".join(text_buf)}</p>')
                return ''.join(out)

            content = tex_to_html(raw_src)
            file_processed = True
            print(f"[DEBUG] TEX imported, raw={len(raw_tex)} chars, html={len(content)} chars")

        else:
            print(f"[DEBUG] Unsupported file type: {filename}. Supported: .txt, .docx, .pdf, .tex")

    except Exception as e:
        print(f"[DEBUG] Error processing file: {e}")
        import traceback
        traceback.print_exc()

    # Only save and print success if file was actually processed
    if file_processed:
        print("File loaded successfully")
        
        # apply sort so new view respects the selected order
        if sort_by:
            citations.sort(
                key=lambda c: (getattr(c, sort_by) or "").lower(),
                reverse=(sort_dir == 'desc')
            )

        save_data()

    # Post-Redirect-Get: redirect to index to avoid reload re-submitting the POST
    return redirect(url_for('index'))


# ---------------- ADD CITATION ----------------
@app.route('/add_citation', methods=['POST'])
def add_citation():

    global next_cid
    global citations
    global content
    global sort_by
    global sort_dir

    cid = next_cid
    next_cid += 1

    citation = Citation(

        cid=cid,

        author=request.form.get('citation_author'),

        year=request.form.get('citation_year'),

        source_type=request.form.get('citation_type'),

        link=request.form.get('citation_link'),

        title=request.form.get('citation_title'),

        issue=request.form.get('citation_journal'),

        publisher=request.form.get('citation_publisher'),

        volume=request.form.get('citation_volume'),

        pages=request.form.get('citation_pages'),

        doi=request.form.get('citation_doi'),

        style=request.form.get('citation_style') or 'APA7'

    )

    # Check for duplicates based on title and author
    if any(c.title == citation.title and c.author == citation.author for c in citations):
        # Don't add duplicate
        pass
    else:
        citations.append(citation)
        print("Citation added:", citation.author)

    # reapply sort
    if sort_by:
        citations.sort(
            key=lambda c: (getattr(c, sort_by) or "").lower(),
            reverse=(sort_dir == 'desc')
        )

    save_data()

    # Post-Redirect-Get: redirect to index to avoid reload re-submitting the POST
    return redirect(url_for('index'))


# ---------------- DELETE CITATION ----------------
@app.route('/delete_citation', methods=['POST'])
def delete_citation():

    global citations
    global content
    global sort_dir
    global sort_by
    # route should handle both form-encoded and JSON bodies
    data = request.get_json(silent=True)
    if data is None:
        # fall back to form data
        cid = request.form.get("citation_id")
    else:
        cid = data.get("citation_id")

    print("DELETE ROUTE HIT, cid =", cid)

    if cid is None:
        return jsonify({"success": False, "error": "no cid"})

    try:
        cid = int(cid)
    except ValueError:
        return jsonify({"success": False, "error": "invalid cid"})

    citation = next((c for c in citations if c.cid == cid), None)
    if citation is None:
        return jsonify({"success": False, "error": "citation not found"}), 404

    content = citation.onDeleteCitation(content)
    print("My new content:", content)
    # remove the citation
    citations = [c for c in citations if c.cid != cid]

    # renumber remaining citations so cids are contiguous starting at 1
    for i, c in enumerate(citations):
        c.cid = i + 1

    # update next_cid so new citations continue after the last current id
    global next_cid
    next_cid = len(citations) + 1

    print("Remaining citations:", len(citations))

    save_data()

    #return render_template("page.html",citations=citations,content=content,sort_by=sort_by,sort_dir=sort_dir)
    return jsonify({"success": True})


# ---------------- UPDATE CITATION STYLE ----------------
@app.route('/update_citation_style', methods=['POST'])
def update_citation_style():
    """Change the style of a single citation."""
    global content
    global sort_by
    global sort_dir
    global citations

    data = request.get_json()
    cid = data.get('citation_id')
    new_style = data.get('style')
    if cid is None or new_style is None:
        return jsonify({"success": False, "error": "missing parameters"}), 400
    try:
        cid = int(cid)
    except ValueError:
        return jsonify({"success": False, "error": "invalid cid"}), 400

    citation = next((c for c in citations if c.cid == cid), None)
    if citation is None:
        return jsonify({"success": False, "error": "not found"}), 404

    content = citation.onUpdateStyle(new_style,content)
    print(f"Citation {cid} style changed to {citation.style}")
    
    save_data()
    
    return jsonify({"success": True})


# ---------------- DUPLICATE CITATION ----------------
@app.route('/duplicate_citation', methods=['POST'])
def duplicate_citation():
    global next_cid, citations, sort_by, sort_dir

    data = request.get_json()
    cid = data.get('citation_id')
    if cid is None:
        return jsonify({"success": False, "error": "no cid"}), 400
    try:
        cid = int(cid)
    except ValueError:
        return jsonify({"success": False, "error": "invalid cid"}), 400

    citation = next((c for c in citations if c.cid == cid), None)
    if citation is None:
        return jsonify({"success": False, "error": "citation not found"}), 404

    # create a copy
    new_citation = Citation(
        cid=next_cid,
        author=citation.author,
        year=citation.year,
        source_type=citation.source_type,
        link=citation.link,
        pages=citation.pages,
        volume=citation.volume,
        issue=citation.issue,
        publisher=citation.publisher,
        doi=citation.doi,
        title=citation.title,
        style=citation.style
    )
    next_cid += 1
    citations.append(new_citation)

    # apply sort
    if sort_by:
        citations.sort(
            key=lambda c: (getattr(c, sort_by) or "").lower(),
            reverse=(sort_dir == 'desc')
        )

    save_data()

    return jsonify({"success": True})


# ---------------- SORT CITATIONS ----------------
@app.route('/sort_citations', methods=['POST'])
def sort_citations():
    """Sort the global citations list by a specified attribute and direction."""

    global sort_by, sort_dir, content, citations

    valid_keys = ['author', 'title', 'publisher', 'pages']
    key = request.form.get('sort_by')
    direction = request.form.get('sort_dir', 'asc')
    if direction not in ['asc', 'desc']:
        direction = 'asc'

    if key not in valid_keys:
        return render_template(
            "page.html",
            citations=citations,
            content=content,
            sort_by=sort_by,
            sort_dir=sort_dir
        )

    sort_by = key
    sort_dir = direction

    citations.sort(
        key=lambda c: (getattr(c, sort_by) or "").lower(),
        reverse=(sort_dir == 'desc')
    )
    print(f"Citations sorted by {sort_by} ({sort_dir})")

    save_data()

    # use Post-Redirect-Get so the selected sort persists without leaving the page in a POST state
    return render_template(
        "page.html",
        citations=citations,
        content=content,
        sort_by=sort_by,
        sort_dir=sort_dir
    )



# ---------------- SAVE TEXT ----------------
@app.route('/save-data', methods=['POST'])
def save_text():

    global content

    data = request.get_json()

    content = data.get('content')

    save_data()

    return jsonify({"success": True})


# ---------------- CLEAR / NEW DOCUMENT ----------------
@app.route('/clear', methods=['POST'])
def clear_document():
    global content, citations, next_cid, sort_by, sort_dir, raw_tex
    content = ''
    citations = []
    next_cid = 1
    sort_by = None
    sort_dir = 'asc'
    raw_tex = ''
    save_data()
    return redirect(url_for('index'))


# ---------------- INSERT CITATION ----------------
@app.route('/insert-citation', methods=['POST'])
def insert_citation():
    global citations

    data = request.get_json() or {}

    try:
        cid = int(data.get('citation_id'))
    except (TypeError, ValueError):
        return jsonify({"error": "invalid citation_id"}), 400

    cursor = data.get('cursor')
    try:
        cursor = int(cursor) if cursor is not None else 0
    except (TypeError, ValueError):
        cursor = 0

    # locate citation by its stable cid value
    citation = next((c for c in citations if c.cid == cid), None)
    if citation is None:
        return jsonify({"error": "citation not found"}), 404
    
    for c in citations:
        c.onNewInsert(cursor)

    inline = citation.inline()  # uses citation.style by default

    print("Inline citation created:", inline)

    save_data()

    return jsonify({
        "citation": inline
    })


# ---------------- EXPORT ----------------
@app.route('/export', methods=['POST'])
def export():
    global raw_tex
    import re
    from html.parser import HTMLParser
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from reportlab.platypus import SimpleDocTemplate, Paragraph as RLParagraph, Spacer, HRFlowable
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors as rl_colors

    title = request.form.get('title', 'document')
    doc_content = request.form.get('content') or content
    fmt = request.form.get('format', 'txt')
    refs_only = request.form.get('refs_only') == '1'

    # If exporting references only, build a simple HTML list of all citations
    if refs_only:
        ref_items = ''.join(
            f'<p>{i+1}. {c.reference(c.style, i+1)}</p>'
            for i, c in enumerate(citations)
        )
        doc_content = f'<h2>References</h2>{ref_items}'
        title = title + '_references'

    # ── colour helper ────────────────────────────────────────────────────────
    def parse_rgb(s):
        m = re.match(r'rgb\((\d+),\s*(\d+),\s*(\d+)\)', s.strip())
        if m:
            return int(m.group(1)), int(m.group(2)), int(m.group(3))
        m = re.match(r'#([0-9a-fA-F]{6})', s.strip())
        if m:
            h = m.group(1)
            return int(h[0:2],16), int(h[2:4],16), int(h[4:6],16)
        return None

    # ── shared event-driven parser base ──────────────────────────────────────
    # We build a simple tag-stack so every handler knows its nesting context.

    class StackParser(HTMLParser):
        """Base: maintains a stack of (tag, attrs_dict) pairs."""
        def __init__(self):
            super().__init__(convert_charrefs=True)
            self.stack = []   # list of (tag, attrs_dict)

        def handle_starttag(self, tag, attrs):
            self.stack.append((tag, dict(attrs)))
            self.on_open(tag, dict(attrs))

        def handle_endtag(self, tag):
            self.on_close(tag)
            # pop the most recent matching entry
            for i in range(len(self.stack)-1, -1, -1):
                if self.stack[i][0] == tag:
                    self.stack.pop(i)
                    break

        def handle_data(self, data):
            self.on_data(data)

        def on_open(self, tag, attrs): pass
        def on_close(self, tag): pass
        def on_data(self, data): pass

        # helpers
        def in_tag(self, *tags):
            return any(t in tags for t, _ in self.stack)

        def current_style(self):
            """Merge all inline style strings in the current stack."""
            merged = {}
            for _, attrs in self.stack:
                for part in attrs.get('style', '').split(';'):
                    if ':' in part:
                        k, _, v = part.partition(':')
                        merged[k.strip()] = v.strip()
            return merged

        def current_attrs(self, tag):
            for t, a in reversed(self.stack):
                if t == tag:
                    return a
            return {}

    # ── TXT ──────────────────────────────────────────────────────────────────

    class PlainTextParser(StackParser):
        BLOCK = {'p','div','h1','h2','h3','h4','li','br','tr'}
        def __init__(self):
            super().__init__()
            self.parts = []
            self._prev_was_block = False
        def on_open(self, tag, attrs):
            if tag in self.BLOCK:
                # Add proper sentence/paragraph spacing
                if self._prev_was_block:
                    self.parts.append('\n\n')  # Double newline between paragraphs
                else:
                    self.parts.append('\n')
                self._prev_was_block = True
        def on_close(self, tag):
            if tag in self.BLOCK:
                self.parts.append('\n')
                self._prev_was_block = False
        def on_data(self, data):
            self.parts.append(data)
        def result(self):
            # Clean up excessive newlines but preserve paragraph spacing
            text = ''.join(self.parts)
            # Replace 3+ newlines with double newline (paragraph break)
            text = re.sub(r'\n{3,}', '\n\n', text)
            return text.strip()

    # ── DOCX ─────────────────────────────────────────────────────────────────

    class DocxParser(StackParser):
        HEADING = {'h1':1,'h2':2,'h3':3,'h4':4}

        def __init__(self, doc):
            super().__init__()
            self.doc = doc
            self.para = None      # current python-docx paragraph
            self._pending = []    # text chunks waiting for a para

        def _ensure_para(self):
            if self.para is None:
                self.para = self.doc.add_paragraph()

        def on_open(self, tag, attrs):
            if tag in self.HEADING:
                self.para = self.doc.add_heading('', level=self.HEADING[tag])
            elif tag == 'p':
                self.para = self.doc.add_paragraph()
                # alignment
                style = attrs.get('style', '')
                if 'text-align: center' in style or 'text-align:center' in style:
                    self.para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif 'text-align: right' in style or 'text-align:right' in style:
                    self.para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif 'text-align: justify' in style or 'text-align:justify' in style:
                    self.para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            elif tag == 'br':
                self._ensure_para()
                self.para.add_run('\n')

        def on_close(self, tag):
            if tag in ('p',) or tag in self.HEADING:
                self.para = None

        def on_data(self, data):
            if not data:
                return
            self._ensure_para()
            run = self.para.add_run(data)
            # bold / italic / underline / strike
            run.bold      = self.in_tag('strong','b')
            run.italic    = self.in_tag('em','i')
            run.underline = self.in_tag('u')
            run.font.strike = self.in_tag('s','strike')
            # inline styles from stack
            style = self.current_style()
            if 'color' in style:
                rgb = parse_rgb(style['color'])
                if rgb:
                    run.font.color.rgb = RGBColor(*rgb)
            if 'font-size' in style:
                sv = style['font-size']
                try:
                    if sv.endswith('pt'):
                        run.font.size = Pt(float(sv[:-2]))
                    elif sv.endswith('px'):
                        run.font.size = Pt(float(sv[:-2]) * 0.75)
                except ValueError:
                    pass
            # Quill size classes on <span class="ql-size-*">
            for _, attrs in self.stack:
                for cls in (attrs.get('class') or '').split():
                    if cls == 'ql-size-small':  run.font.size = Pt(8)
                    if cls == 'ql-size-large':  run.font.size = Pt(16)
                    if cls == 'ql-size-huge':   run.font.size = Pt(24)

    def build_docx(html, include_refs=True):
        doc = Document()
        doc.add_heading(title, level=1)
        p = DocxParser(doc)
        p.feed(html)
        if include_refs and citations:
            doc.add_heading('References', level=2)
            for i, c in enumerate(citations):
                doc.add_paragraph(c.reference(c.style, i + 1))
        return doc

    # ── PDF ──────────────────────────────────────────────────────────────────
    # ReportLab Paragraph accepts a limited XML subset:
    # <b>, <i>, <u>, <font color="#rrggbb" size="N">, <br/>
    # We convert our HTML to that markup string per paragraph.

    class RLMarkupParser(StackParser):
        """Convert a single HTML block element's content to RL markup."""
        def __init__(self):
            super().__init__()
            self.parts = []

        def on_open(self, tag, attrs):
            if tag in ('strong','b'):  self.parts.append('<b>')
            elif tag in ('em','i'):    self.parts.append('<i>')
            elif tag == 'u':           self.parts.append('<u>')
            elif tag == 'br':          self.parts.append('<br/>')
            elif tag == 'span':
                style = attrs.get('style','')
                opens = ''
                for part in style.split(';'):
                    part = part.strip()
                    if part.startswith('color:'):
                        rgb = parse_rgb(part[6:])
                        if rgb:
                            opens += '<font color="#{:02x}{:02x}{:02x}">'.format(*rgb)
                    elif part.startswith('font-size:'):
                        sv = part[10:].strip()
                        try:
                            if sv.endswith('pt'):
                                opens += f'<font size="{sv[:-2]}">'
                            elif sv.endswith('px'):
                                opens += f'<font size="{float(sv[:-2])*0.75:.1f}">'
                        except ValueError:
                            pass
                self.parts.append(opens)

        def on_close(self, tag):
            if tag in ('strong','b'):  self.parts.append('</b>')
            elif tag in ('em','i'):    self.parts.append('</i>')
            elif tag == 'u':           self.parts.append('</u>')
            elif tag == 'span':
                # close as many <font> tags as we opened
                style = self.current_attrs('span').get('style','')
                closes = style.count('color:') + style.count('font-size:')
                self.parts.append('</font>' * closes)

        def on_data(self, data):
            self.parts.append(
                data.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;')
            )

        def result(self):
            return ''.join(self.parts)

    class PDFBlockParser(StackParser):
        """Walk top-level blocks and emit ReportLab flowables."""
        HEADING  = {'h1', 'h2', 'h3', 'h4'}
        BLOCK    = {'p', 'div'} | HEADING
        LIST_TAG = {'ul', 'ol'}

        def __init__(self, rl_styles, normal_style):
            super().__init__()
            self.elements = []
            self._rl_styles   = rl_styles
            self._normal      = normal_style
            self._buf         = []
            self._block_tag   = None
            self._block_attrs = {}
            self._depth       = 0
            # list state
            self._in_list     = False
            self._list_type   = None   # 'ul' or 'ol'
            self._item_buf    = []
            self._item_num    = 0
            self._in_item     = False

        def on_open(self, tag, attrs):
            if self._in_list:
                if tag == 'li':
                    self._in_item = True
                    self._item_buf = []
                    self._item_num += 1
                elif self._in_item:
                    self._depth += 1
                    attr_str = ''.join(f' {k}="{v}"' for k, v in attrs.items())
                    self._item_buf.append(f'<{tag}{attr_str}>')
                return
            if tag in self.LIST_TAG and self._depth == 0:
                self._in_list  = True
                self._list_type = tag
                self._item_num  = 0
                return
            if tag in self.BLOCK and self._depth == 0:
                self._block_tag   = tag
                self._block_attrs = attrs
                self._buf         = []
                self._depth       = 1
            elif self._depth > 0:
                self._depth += 1
                attr_str = ''.join(f' {k}="{v}"' for k, v in attrs.items())
                self._buf.append(f'<{tag}{attr_str}>')

        def on_close(self, tag):
            if self._in_list:
                if tag in self.LIST_TAG:
                    self._in_list = False
                    self._list_type = None
                elif tag == 'li':
                    self._flush_item()
                    self._in_item = False
                elif self._in_item:
                    self._depth -= 1
                    self._item_buf.append(f'</{tag}>')
                return
            if self._depth == 0:
                return
            if tag in self.BLOCK and self._depth == 1:
                self._flush()
                self._depth     = 0
                self._block_tag = None
            else:
                self._depth -= 1
                self._buf.append(f'</{tag}>')

        def on_data(self, data):
            safe = data.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;')
            if self._in_list and self._in_item:
                self._item_buf.append(safe)
            elif self._depth > 0:
                self._buf.append(safe)

        def _flush_item(self):
            inner = ''.join(self._item_buf)
            mp = RLMarkupParser()
            mp.feed(inner)
            markup = mp.result().strip()
            if not markup:
                return
            bullet = f'{self._item_num}.' if self._list_type == 'ol' else '\u2022'
            item_style = _unique_style(
                'li', parent=self._normal,
                leftIndent=20, firstLineIndent=-12, spaceAfter=2
            )
            try:
                self.elements.append(RLParagraph(f'{bullet} {markup}', item_style))
            except Exception:
                plain = re.sub(r'<[^>]+>', '', markup)
                self.elements.append(RLParagraph(f'{bullet} {plain}', item_style))

        def _flush(self):
            inner_html = ''.join(self._buf)
            mp = RLMarkupParser()
            mp.feed(inner_html)
            markup = mp.result().strip()

            if not markup or markup == '<br/>':
                self.elements.append(Spacer(1, 0.15 * inch))
                return

            tag = self._block_tag
            if tag == 'h1':
                st = self._rl_styles['h1']
            elif tag == 'h2':
                st = self._rl_styles['h2']
            elif tag in ('h3','h4'):
                st = self._rl_styles['h3']
            else:
                st = _unique_style('pl', parent=self._normal)
                style_str = self._block_attrs.get('style','')
                if 'text-align: center' in style_str or 'text-align:center' in style_str:
                    st.alignment = 1
                elif 'text-align: right' in style_str or 'text-align:right' in style_str:
                    st.alignment = 2

            try:
                self.elements.append(RLParagraph(markup, st))
            except Exception:
                plain = re.sub(r'<[^>]+>', '', markup)
                self.elements.append(RLParagraph(plain, st))

    # Use a module-level counter so ParagraphStyle names are always unique
    _pdf_style_counter = [0]

    def _unique_style(name, **kwargs):
        _pdf_style_counter[0] += 1
        return ParagraphStyle(f'{name}_{_pdf_style_counter[0]}', **kwargs)

    def build_pdf(html, include_refs=True):
        buf      = io.BytesIO()
        rl_styles_base = getSampleStyleSheet()
        normal   = _unique_style('Body',  parent=rl_styles_base['Normal'],
                                  fontSize=11, leading=15, spaceAfter=4)
        rl_styles = {
            'h1': _unique_style('H1', parent=rl_styles_base['Heading1'], fontSize=18, spaceAfter=10),
            'h2': _unique_style('H2', parent=rl_styles_base['Heading2'], fontSize=15, spaceAfter=8),
            'h3': _unique_style('H3', parent=rl_styles_base['Heading3'], fontSize=13, spaceAfter=6),
        }
        ref_style = _unique_style('Ref', parent=normal, fontSize=10, leftIndent=20)

        doc = SimpleDocTemplate(buf, pagesize=letter,
                                leftMargin=inch, rightMargin=inch,
                                topMargin=inch, bottomMargin=inch)
        elements = []
        title_safe = title.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;')
        elements.append(RLParagraph(title_safe, rl_styles['h1']))
        elements.append(Spacer(1, 0.1 * inch))

        bp = PDFBlockParser(rl_styles, normal)
        bp.feed(html)
        elements.extend(bp.elements)

        if include_refs and citations:
            elements.append(Spacer(1, 0.2 * inch))
            elements.append(HRFlowable(width='100%', thickness=1, color=rl_colors.grey))
            elements.append(Spacer(1, 0.1 * inch))
            elements.append(RLParagraph('References', rl_styles['h2']))
            for i, c in enumerate(citations):
                ref = c.reference(c.style, i+1)
                ref_safe = ref.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;')
                elements.append(RLParagraph(ref_safe, ref_style))

        doc.build(elements)
        buf.seek(0)
        return buf

    # ── dispatch ─────────────────────────────────────────────────────────────

    if fmt == 'txt':
        p = PlainTextParser()
        p.feed(doc_content)
        plain = p.result()
        if not refs_only:
            ref_lines = '\n'.join(c.reference(c.style, i+1) for i, c in enumerate(citations))
            plain += ('\n\nReferences:\n' + ref_lines if ref_lines else '')
        buf = io.BytesIO(plain.encode('utf-8'))
        return send_file(buf, as_attachment=True, download_name=title + '.txt',
                         mimetype='text/plain')

    elif fmt == 'docx':
        doc = build_docx(doc_content, include_refs=not refs_only)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return send_file(buf, as_attachment=True, download_name=title + '.docx',
                         mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    elif fmt == 'pdf':
        try:
            buf = build_pdf(doc_content, include_refs=not refs_only)
            return send_file(buf, as_attachment=True, download_name=title + '.pdf',
                             mimetype='application/pdf')
        except Exception as e:
            print(f'[DEBUG] PDF export error: {e}')
            import traceback; traceback.print_exc()
            return jsonify({'error': 'PDF generation failed', 'detail': str(e)}), 500

    elif fmt == 'tex':
        import re as _re3, html as _html3
        def html_to_latex(h):
            class TexP(HTMLParser):
                def __init__(self):
                    super().__init__(convert_charrefs=True)
                    self.out = []
                    self.stack = []
                def handle_starttag(self, tag, attrs):
                    self.stack.append(tag)
                    if tag in ('strong','b'):   self.out.append('\\textbf{')
                    elif tag in ('em','i'):     self.out.append('\\textit{')
                    elif tag == 'u':            self.out.append('\\underline{')
                    elif tag == 'code':         self.out.append('\\texttt{')
                    elif tag == 'blockquote':   self.out.append('\n\\begin{quote}\n')
                    # heading levels mirror what tex_to_html produces on import:
                    # \chapter->\h1, \section->\h2, \subsection->\h3, \subsubsection->\h4
                    elif tag == 'h1':           self.out.append('\n\\chapter{')
                    elif tag == 'h2':           self.out.append('\n\\section{')
                    elif tag == 'h3':           self.out.append('\n\\subsection{')
                    elif tag == 'h4':           self.out.append('\n\\subsubsection{')
                    elif tag == 'h5':           self.out.append('\n\\paragraph{')
                    elif tag == 'li':           self.out.append('\n  \\item ')
                    elif tag == 'ul':           self.out.append('\n\\begin{itemize}\n')
                    elif tag == 'ol':           self.out.append('\n\\begin{enumerate}\n')
                    elif tag == 'br':           self.out.append('\n')
                def handle_endtag(self, tag):
                    if self.stack and self.stack[-1] == tag:
                        self.stack.pop()
                    if tag in ('strong','b','em','i','u','code','h1','h2','h3','h4','h5'):
                        self.out.append('}')
                    elif tag == 'p':
                        self.out.append('\n\n')
                    elif tag == 'blockquote':
                        self.out.append('\n\\end{quote}\n')
                    elif tag == 'ul':
                        self.out.append('\n\\end{itemize}\n')
                    elif tag == 'ol':
                        self.out.append('\n\\end{enumerate}\n')
                def handle_data(self, data):
                    ESC = [('\\','\\textbackslash{}'),('&','\\&'),('%','\\%'),
                           ('$','\\$'),('#','\\#'),('_','\\_'),('{','\\{'),
                           ('}','\\}'),('~','\\textasciitilde{}')]
                    for ch, esc in ESC:
                        data = data.replace(ch, esc)
                    self.out.append(data)
                def result(self):
                    import re as _rer
                    return _rer.sub(r'\n{3,}', '\n\n', ''.join(self.out)).strip()
            p = TexP(); p.feed(h); return p.result()

        bib = ''
        if citations:
            bib = '\n\\begin{thebibliography}{99}\n'
            for i, c in enumerate(citations):
                bib += f'  \\bibitem{{ref{i+1}}} {c.reference(c.style, i+1)}\n'
            bib += '\\end{thebibliography}\n'

        if raw_tex:
            new_body = html_to_latex(doc_content)
            if _re3.search(r'\\begin\{document\}', raw_tex):
                out_tex = _re3.sub(
                    r'(\\begin\{document\})(.*?)(\\end\{document\})',
                    lambda m: m.group(1) + '\n' + new_body + '\n' + bib + '\n' + m.group(3),
                    raw_tex, flags=_re3.DOTALL
                )
            else:
                out_tex = raw_tex + '\n\n' + new_body + '\n' + bib
        else:
            body = html_to_latex(doc_content)
            import html as _h4
            out_tex = (
                '\\documentclass[12pt]{article}\n'
                '\\usepackage[utf8]{inputenc}\n'
                '\\usepackage{hyperref}\n'
                f'\\title{{{_h4.escape(title)}}}\n'
                '\\date{}\n'
                '\\begin{document}\n'
                '\\maketitle\n\n'
                + body + '\n' + bib +
                '\n\\end{document}\n'
            )
        buf = io.BytesIO(out_tex.encode('utf-8'))
        return send_file(buf, as_attachment=True, download_name=title + '.tex',
                         mimetype='text/x-tex')


    return jsonify({'error': 'unknown format'}), 400


# ---------------- CITATION PREVIEW ----------------
@app.route('/citation_preview', methods=['POST'])
def citation_preview():
    data = request.get_json(silent=True) or {}
    try:
        cid = int(data.get('citation_id'))
    except (TypeError, ValueError):
        return jsonify({'error': 'invalid cid'}), 400
    c = next((c for c in citations if c.cid == cid), None)
    if not c:
        return jsonify({'error': 'not found'}), 404
    return jsonify({'preview': c.reference(c.style, cid)})


# ---------------- DOI LOOKUP ----------------
@app.route('/doi_lookup', methods=['POST'])
def doi_lookup():
    import urllib.request, urllib.error, json as _json, urllib.parse
    data = request.get_json(silent=True) or {}
    doi = (data.get('doi') or '').strip()
    doi = doi.replace('https://doi.org/','').replace('http://doi.org/','')
    if not doi:
        return jsonify({'error': 'No DOI provided'}), 400
    url = 'https://api.crossref.org/works/' + urllib.parse.quote(doi, safe='')
    try:
        req = urllib.request.Request(url, headers={'User-Agent': 'Docuswitcher/1.0'})
        with urllib.request.urlopen(req, timeout=8) as resp:
            body = _json.loads(resp.read().decode('utf-8'))
        msg = body.get('message', {})
        authors = []
        for a in msg.get('author', []):
            family = a.get('family','')
            given  = a.get('given','')
            authors.append(f'{family}, {given}' if given else family)
        date_parts = msg.get('published', msg.get('published-print', msg.get('published-online', {})))
        year = str(date_parts.get('date-parts',[['']])[0][0]) if date_parts else ''
        container = ''
        for key in ('container-title','publisher'):
            val = msg.get(key)
            if val:
                container = val[0] if isinstance(val, list) else val
                break
        title_val = msg.get('title',[''])[0] if msg.get('title') else ''
        return jsonify({
            'author':    '; '.join(authors),
            'year':      year,
            'title':     title_val,
            'publisher': container,
            'volume':    msg.get('volume',''),
            'issue':     msg.get('issue',''),
            'pages':     msg.get('page',''),
            'doi':       doi,
        })
    except urllib.error.HTTPError as e:
        return jsonify({'error': f'DOI not found (HTTP {e.code})'}), 404
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ---------------- WORD COUNT ----------------
@app.route('/word_count', methods=['POST'])
def word_count():
    import re as _re
    data = request.get_json(silent=True) or {}
    html_content = data.get('content', '')
    plain = _re.sub(r'<[^>]+>', ' ', html_content)
    plain = plain.replace('&nbsp;',' ').replace('&amp;','&').replace('&lt;','<').replace('&gt;','>')
    words = [w for w in _re.split(r'\s+', plain.strip()) if w]
    chars_no_spaces = len(_re.sub(r'\s','',plain))
    return jsonify({'words': len(words), 'chars': len(plain.strip()), 'chars_no_spaces': chars_no_spaces})


if __name__ == '__main__':
    app.run(debug=True)
